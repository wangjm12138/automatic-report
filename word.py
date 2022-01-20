import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.styles.styles import Styles
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
###http://officeopenxml.com/WPtableBorders.php
class Word:
    def __init__(self, Text, feature_defect, defect_feature, loghandle):
        self.Text = Text
        self.feature_defect = feature_defect
        self.defect_feature = defect_feature
        self.loghandle = loghandle
        self.document = Document()

    def add_hyperlink(self,paragraph, text, url):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run ()
        r._r.append (hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return hyperlink

    def paragraphrun_style(self, paragraphrun, bold=True, size=None, name=None):
        paragraphrun.font.bold = bold
        if size is None:
            paragraphrun.font.size = Pt(14)
        else:
            paragraphrun.font.size = Pt(int(size))

        if name is None:
            paragraphrun.font.name = "Arial"
        else:
            paragraphrun.font.name = name
        paragraphrun.font.color.rgb = RGBColor(0,0,0)

    def generate_first_page(self):
        default_paragrah =  self.document.add_paragraph()
        default_paragrah.style.font.size = Pt(11)
        default_paragrah.style.font.name = "Calibri(Body)"
        self.document.add_paragraph()

        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        title = paragraph.add_run(self.Text["Title"]["name"]+"\n")
        self.paragraphrun_style(title, size=20)

        summary = paragraph.add_run(self.Text["Title"]["summary"]+"\n")
        self.paragraphrun_style(summary, size=12)

        hw = paragraph.add_run(self.Text["Title"]["hw"]+"\n")
        self.paragraphrun_style(hw, size=12)

        line = paragraph.add_run(self.Text["Title"]["line"]+"\n")
        self.paragraphrun_style(line, bold=False, size=12)
        version = paragraph.add_run(self.Text["Title"]["version"]+"\n")
        self.paragraphrun_style(version, bold=False, size=16)

        self.document.add_page_break()


    def feature_defect_content(self, table):
        #for data in self.Data:
        index = 1
        for f_k, f_v in self.feature_defect.items():
            cell = table.cell(index, 0)
            p = cell.add_paragraph()
            self.add_hyperlink(p, f_k, f_v["feature_link"])
            cell = table.cell(index, 1)
            p = cell.add_paragraph(f_v["feature_name"])
            cell = table.cell(index, 2)
            for d_k, d_v in f_v["defectlist"].items():
                #cell = table.cell(index, 2)
                if len(d_v["defect_link"]) > 0:
                    p = cell.add_paragraph()
                    defect_link = d_v["defect_link"]
                    jira_num = defect_link.split("/")[-1]
                    self.add_hyperlink(p, jira_num, defect_link)
            index = index +1

    def defect_feature_content(self, table):
        #for data in self.Data:
        index = 1
        for d_k, d_v in self.defect_feature.items():
            cell = table.cell(index, 0)
            p = cell.add_paragraph()
            defect_link = d_v["defect_link"]
            jira_num = defect_link.split("/")[-1]
            self.add_hyperlink(p, jira_num, defect_link)
            cell = table.cell(index, 1)
            p = cell.add_paragraph(d_v["defect_name"])
            cell = table.cell(index, 2)
            p = cell.add_paragraph(d_v["priority"])
            cell = table.cell(index, 3)
            for f_k, f_v in d_v["featurelist"].items():
                #cell = table.cell(index, 2)
                p = cell.add_paragraph()
                feature_link = f_v["feature_link"]
                self.add_hyperlink(p, f_k, feature_link)
            index = index +1

    def heading1_content(self):
        content1 = "Legend: Green for no blockers. Yellow for max two majors. Red for more than two majors.\n"
        table1 = self.document.add_table(rows=1, cols=2, style='Table Grid')
        cell = table1.cell(0,0)
        cell.text = "Number of features"
        cell = table1.cell(0,1)
        cell.text = str(len(self.feature_defect))

        paragraph = self.document.add_paragraph()
        table2 = self.document.add_table(rows=1, cols=2, style='Table Grid')
        cell = table2.cell(0,0)
        cell.text = "Number of known defects"
        cell = table2.cell(0,1)
        cell.text = str(len(self.defect_feature))

        paragraph = self.document.add_paragraph()
        table3 = self.document.add_table(rows=1, cols=2, style='Table Grid')
        cell = table3.cell(0,0)
        cell.text = "Traffic Light Indicator"
        cell = table3.cell(0,1)
        cell.text = "major issues"

        paragraph = self.document.add_paragraph()
        paragraph.add_run(content1)

    def heading2_1_content(self):
        content1 = "This Test Report documents a summary of the features and related known defects supported by this platform release.\nTo give a simple and quick overview, the most important parts of this document are: \n"
        content2 = "Traffic light."
        content3 = "Indicates if this Platform Release is working as expected on the agreed Hardware Configurations."
        content4 = "Open errors."
        content5 = "List of open errors, the severity and the Features affected by them"
        content6 = "Supported Features."
        content7 = "List of supported features and the errors found during testing."
        content8 = "For further details on test results or specific errors, links to Jama and Jira are available to enable the reader for more in-depth investigations."
        paragraph = self.document.add_paragraph()
        paragraph.add_run(content1)
        traffic_light = paragraph.add_run(content2)
        self.paragraphrun_style(traffic_light, bold=True, size=11, name="Calibri(Body)")
        paragraph.add_run(content3+"\n")
        open_errors = paragraph.add_run(content4)
        self.paragraphrun_style(open_errors, bold=True, size=11, name="Calibri(Body)")
        paragraph.add_run(content5+"\n")
        supported_features = paragraph.add_run(content6)
        self.paragraphrun_style(supported_features, bold=True, size=11, name="Calibri(Body)")
        paragraph.add_run(content7+"\n")
        paragraph.add_run(content8)
        #self.paragraphrun_style(purpose, bold=False, size=11, name="Calibri(Body)")
    def heading2_2_content(self):
        content1 = "The Platform Release has been tested on agreed hardware configurations. This must be considered carefully when trying the Platform Release on a new device. Small differences may have an impact on the expected functionality.\n"
        content2 = "Even though this document lists the features supported by the Platform Release, there may be a limit on which features can be simultaneously active.\n"
        content3 = "Test of features has been done to an extent which seems fair from test point of view. There may be combinations / configurations which have not been considered.\n"
        content4 = "Traffic light is an overall consideration of the Platform Release. It is not a guarantee that every feature will work smoothly. The list of known defects must be considered carefully.\n"
        paragraph = self.document.add_paragraph()
        paragraph.add_run(content1)
        paragraph.add_run(content2)
        paragraph.add_run(content3)
        paragraph.add_run(content4)

    def heading2_3_content(self):
        content1 = "Platform Cougar is the platform to be used for testing features. Future projects are supposed to inherit the results from Platform Cougar. Platform Cougar consists of 4 different hardware Blueprints. Rio, Maple Spiderman, Darwin, and Hulk. Test Items will indicate what hardware to be used for testing. In situation, where the indicated platform will not cover the feature, the features can be tested on other hardware, e.g., device R&D PCB or device form factor where the feature is supported.\n"
        paragraph = self.document.add_paragraph()
        paragraph.add_run(content1)

    def heading2_4_content(self):
        content1 = "Test levels are a mix of “End to end tests”, Call control tests (IOP tests) and GN protocol tests.\n"
        content2 = " As soon as a test has been created. It will become a regression test on future PI-releases Feature details for tests in scope for PI-13 are presented in the:\n"
        paragraph = self.document.add_paragraph()
        paragraph.add_run(content1)
        paragraph.add_run(content2)

    def heading2_5_content(self):
        content1 = "5 different hardware blueprints represents Platform Feature Testing.\n"
        content2 = "Software for the hardware blueprints are located in the folder:\n"
        content3 = "\\\\dkcphweb13vm\\releases\\platform\\V13\\1.0 \n"
        content4 = "Darwin: release_1.57.3_211007 \n"
        content5 = "Hulk: release_1.59.1_211007 \n"
        content6 = "Maple: release_0.4.2_211007 \n"
        content7 = "Rio: release_4.57.5_211007 \n"
        content8 = "Spiderman: release_2.60.3_211007 \n"
        paragraph = self.document.add_paragraph()
        paragraph.add_run(content1)
        paragraph.add_run(content2)
        paragraph.add_run(content3)
        paragraph.add_run(content4)
        paragraph.add_run(content5)
        paragraph.add_run(content6)
        paragraph.add_run(content7)
        paragraph.add_run(content8)

    def heading3_content(self):
        content1 = "Details of test results for PI-13 are presented in:"
        paragraph = self.document.add_paragraph()
        purpose = paragraph.add_run(content1)


    def heading3_1_content(self):
        content1 = "List of supported features and the known defects."
        paragraph = self.document.add_paragraph()
        purpose = paragraph.add_run(content1)
        rows = len(self.feature_defect)

        table = self.document.add_table(rows=rows+1, cols=3, style='Table Grid')
        cell = table.cell(0,0)
        cell.text = "Feature"
        cell = table.cell(0,1)
        cell.text = "Name"
        cell = table.cell(0,2)
        cell.text = "Known Defects"
        self.feature_defect_content(table)
        #paragraph = cell.add_paragraph()

        #self.add_hyperlink(paragraph, "baidu","https://www.baidu.com")
    def heading4_content(self):
        content1 = "The number of open defects with severity. Each defect can find more details such steps of reproducing, effected HW etc. via the link including the linked feature that affected the issue."
        paragraph = self.document.add_paragraph()
        purpose = paragraph.add_run(content1)
        rows = len(self.defect_feature)

        table = self.document.add_table(rows=rows+1, cols=3, style='Table Grid')
        cell = table.cell(0,0)
        cell.text = "Issue"
        cell = table.cell(0,1)
        cell.text = "Name"
        cell = table.cell(0,2)
        cell.text = "Severity"
        cell = table.cell(0,3)
        cell.text = "Affected Feature"
        self.defect_feature_content(table)

    def generate_second_page(self):
        heading1 = self.document.add_heading(level=1)
        SUMMARY = heading1.add_run(self.Text["Heading"]["heading1"]+"\n")
        self.paragraphrun_style(SUMMARY)
        self.heading1_content()

        heading2 = self.document.add_heading(level=1)
        INTRODUCTION = heading2.add_run(self.Text["Heading"]["heading2"]+"\n")
        self.paragraphrun_style(INTRODUCTION)

        heading2_1 = self.document.add_heading(level=2)
        PURPOSE = heading2_1.add_run(self.Text["Heading"]["heading2.1"]+"\n")
        self.paragraphrun_style(PURPOSE)
        self.heading2_1_content()

        heading2_2 = self.document.add_heading(level=2)
        HOWTOUSE = heading2_2.add_run(self.Text["Heading"]["heading2.2"]+"\n")
        self.paragraphrun_style(HOWTOUSE)
        self.heading2_2_content()

        heading2_3 = self.document.add_heading(level=2)
        OVERVIEW = heading2_3.add_run(self.Text["Heading"]["heading2.3"]+"\n")
        self.paragraphrun_style(OVERVIEW)
        self.heading2_3_content()

        heading2_4 = self.document.add_heading(level=2)
        SCOPE = heading2_4.add_run(self.Text["Heading"]["heading2.4"]+"\n")
        self.paragraphrun_style(SCOPE)
        self.heading2_4_content()

        heading2_5 = self.document.add_heading(level=2)
        SETUP = heading2_5.add_run(self.Text["Heading"]["heading2.5"]+"\n")
        self.paragraphrun_style(SETUP)
        self.heading2_5_content()

        heading3 = self.document.add_heading(level=1)
        TEST_RESULTS = heading3.add_run(self.Text["Heading"]["heading3"]+"\n")
        self.paragraphrun_style(TEST_RESULTS)
        self.heading3_content()


        heading3_1 = self.document.add_heading(level=2)
        SUPPORTED_FEATURES = heading3.add_run(self.Text["Heading"]["heading3.1"]+"\n")
        self.paragraphrun_style(SUPPORTED_FEATURES)
        self.heading3_1_content()

        heading4 = self.document.add_heading(level=1)
        KNOWN_DEFECTS = heading4.add_run(self.Text["Heading"]["heading4"]+"\n")
        self.paragraphrun_style(KNOWN_DEFECTS)
        self.heading4_content()

    def generate_word(self):
        self.generate_first_page()
        self.generate_second_page()
        #paragraph = self.document.add_paragraph(self.Text["Title"]["name"])


# section = document.section[0]
# header = section.header
# paragraph = header.paragraphs[0]
#
# paragraph.text = "\t\tTEMPLATE v0.1"
# Title
# paragraph.style.font.size = Pt(20)
#
# document.add_heading('Document Title', 0)
    def save(self):
        self.document.save('demo.docx')
